"""
Phase 5B : Chunking Intelligent et Classification Chunk-Level
Architecture hybride : Découpage structurel + Qualification LLM légère

Workflow :
1. Chunking structurel (HTML titres, PDF TOC, etc.)
2. Features heuristiques (regex, signaux lexicaux)
3. Qualification LLM (chunks ambigus seulement)
4. Fusion intelligente (heuristiques > LLM)
5. Output : processed_chunks.jsonl
"""

import json
from pathlib import Path
import sys
import logging
import signal
from typing import Dict, List
from tqdm import tqdm
import re
from bs4 import BeautifulSoup

# Ajouter chemins
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root / 'src' / 'utils'))
sys.path.insert(0, str(project_root / 'src' / 'processing'))

from llm_provider import RAGConfig
from json_cleaner import clean_llm_json_response, safe_parse_json
from rgpd_topics import parse_tags, TAG_PROMPT_TABLE

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)
logging.getLogger("httpx").setLevel(logging.WARNING)


class ChunkFeatureExtractor:
    """Extraction features heuristiques (sans LLM)"""
    
    # Patterns
    ARTICLE_RGPD = re.compile(r'article\s+\d+', re.IGNORECASE)
    STEPS = re.compile(r'étape|procédure|1\.|2\.|3\.', re.IGNORECASE)
    TEMPLATE = re.compile(r'modèle|template|formulaire|checklist', re.IGNORECASE)
    SANCTION = re.compile(r'san-\d+|amende|mise en demeure', re.IGNORECASE)
    TECHNICAL = re.compile(r'chiffrement|cryptographie|tls|ssl|https|pseudonymisation|anonymisation|hash|sécurité|mesure technique|authentification|journalisation|accès|contrôle', re.IGNORECASE)
    LEGAL_REF = re.compile(r'règlement\s+\(ue\)|directive|loi\s+n°', re.IGNORECASE)
    
    # Keywords doctrine vs opérationnel
    DOCTRINE_KEYWORDS = ['principe', 'définition', 'notion', 'portée', 'interprétation', 'clarification']
    OPERATIONAL_KEYWORDS = ['vérifier', 'contrôler', 'mettre en place', 'mesure', 'exemple', 'cas pratique']
    
    # Secteurs
    SECTOR_KEYWORDS = {
        "Santé": ["dossier médical", "patient", "hôpital", "soin", "santé"],
        "RH": ["salarié", "recrutement", "paie", "ressources humaines"],
        "Marketing": ["prospection", "cookies", "tracking", "publicité"],
        "Éducation": ["école", "université", "étudiant", "élève"],
    }
    
    @classmethod
    def extract(cls, text: str) -> Dict:
        text_lower = text.lower()
        
        # Scores doctrine vs opérationnel
        doctrine_score = sum(1 for kw in cls.DOCTRINE_KEYWORDS if kw in text_lower)
        operational_score = sum(1 for kw in cls.OPERATIONAL_KEYWORDS if kw in text_lower)
        
        # Secteurs détectés
        sectors = []
        for sector, keywords in cls.SECTOR_KEYWORDS.items():
            if any(kw in text_lower for kw in keywords):
                sectors.append(sector)
        
        return {
            'has_article_rgpd': bool(cls.ARTICLE_RGPD.search(text)),
            'has_legal_ref': bool(cls.LEGAL_REF.search(text)),
            'has_steps': bool(cls.STEPS.search(text)),
            'has_template': bool(cls.TEMPLATE.search(text)),
            'has_sanction': bool(cls.SANCTION.search(text)),
            'has_technical': bool(cls.TECHNICAL.search(text)),
            'has_numbered_list': bool(re.search(r'^\s*\d+\.\s', text, re.MULTILINE)),
            'doctrine_score': doctrine_score,
            'operational_score': operational_score,
            'sectors': sectors,
            'word_count': len(text.split()),
        }


class StructuralChunker:
    """Découpage structurel avec overlap et split sémantique.
    
    Principes :
    - Split aux frontières de phrases (\n\n ou .), pas au milieu des mots
    - Overlap entre chunks consécutifs pour ne pas couper les concepts
    - Heading propagé et préfixé dans le texte pour l'embedding
    - Tiny chunks (<100w) fusionnés avec leur voisin
    """
    
    def __init__(self):
        self.target_size = 400  # mots — taille idéale
        self.min_size = 100     # en dessous → fusionner avec voisin
        self.max_size = 600     # au dessus → split sémantique
        self.overlap = 50       # mots de chevauchement entre chunks consécutifs
    
    def chunk_html(self, file_path: Path) -> List[Dict]:
        """Chunking HTML — cible le contenu principal (region-content)"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f, 'lxml')
            
            # Cibler le bloc de contenu principal (structure CNIL)
            content_block = (
                soup.find(class_='region-content')
                or soup.find('main')
                or soup.find('article')
                or soup.find(class_='field-name-body')
                or soup  # Fallback: page entière
            )
            
            # Nettoyage des éléments non-contenu
            for tag in content_block(['script', 'style', 'nav', 'header', 'footer',
                                       'aside', 'iframe', 'noscript', 'svg']):
                tag.decompose()
            
            # Supprimer menus, breadcrumbs, pagination
            for nav_block in content_block.find_all(class_=lambda c: c and any(
                x in str(c).lower() for x in [
                    'menu-push', 'breadcrumb', 'pager', 'pagination',
                    'nav-', 'share-', 'social', 'cookie', 'back-to-top'
                ]
            )):
                nav_block.decompose()
            
            chunks = []
            current = {'text': '', 'heading': ''}
            
            for elem in content_block.find_all(['h2', 'h3', 'p', 'ul', 'table']):
                if elem.name in ['h2', 'h3']:
                    if current['text']:
                        chunks.append(current)
                    current = {'text': elem.get_text(strip=True), 'heading': elem.get_text(strip=True)}
                elif elem.name == 'table':
                    # ── Détection content-based : table HTML → pipeline LLM ──
                    # Flush le texte courant avant la table
                    if current['text']:
                        chunks.append(current)
                        current = {'text': '', 'heading': current.get('heading', '')}
                    
                    table_chunks = self._convert_html_table(elem, current.get('heading', ''))
                    if table_chunks:
                        chunks.extend(table_chunks)
                else:
                    current['text'] += ' ' + elem.get_text(strip=True)
            
            if current['text']:
                chunks.append(current)
            
            return self._post_process(chunks)
        except:
            return []
    
    def _convert_html_table(self, table_elem, heading: str) -> List[Dict]:
        """Convertit une <table> HTML en chunks via le pipeline LLM tableur.
        
        1. Extraire les lignes (th/td) comme listes de cellules
        2. Appeler _convert_table_rows() (pipeline commun)
        """
        # Extraire les lignes comme listes de cellules (même format que spreadsheets)
        rows = []
        for tr in table_elem.find_all('tr'):
            cells = [c.get_text(strip=True) for c in tr.find_all(['th', 'td'])]
            if any(c for c in cells):
                rows.append(cells)
        
        if len(rows) < 2:
            # Table triviale → texte plat suffit
            flat = table_elem.get_text(strip=True)
            if flat:
                return [{'text': flat, 'heading': heading}]
            return []
        
        return self._convert_table_rows(rows, heading)
    
    def _convert_table_rows(self, rows: List[List[str]], heading: str) -> List[Dict]:
        """Pipeline commun : convertit des lignes de cellules en chunks via LLM.
        
        Utilisé par : HTML <table>, DOCX tables, et potentiellement PDF tables.
        Réutilise le pipeline spreadsheet (zones, split, pipe-text, LLM Nemo).
        """
        # Initialiser LLM si pas encore fait
        if not hasattr(self, '_llm'):
            try:
                config = RAGConfig()
                self._llm = config.llm_provider
            except Exception as e:
                logger.warning(f"LLM indisponible pour conversion tableau: {e}")
                self._llm = None
        
        context = heading or 'Tableau'
        zones = self._segment_sheet_zones(rows)
        
        chunks = []
        for zone_rows in zones:
            sub_zones = self._split_large_zone(zone_rows)
            for sub_zone in sub_zones:
                raw_table_text = self._zone_to_table_text(sub_zone, context)
                if not raw_table_text or len(raw_table_text.split()) < 10:
                    continue
                
                llm_result = self._llm_convert_table(raw_table_text, context)
                natural_text = llm_result['text']
                tags = llm_result.get('tags', [])
                
                if not natural_text or len(natural_text.split()) < 15:
                    continue
                
                sub_chunks = self._split_long_text(natural_text, heading, tags=tags, max_words=500)
                chunks.extend(sub_chunks)
        
        return chunks
    
    def _convert_docx_table(self, rows: List[List[str]], heading: str) -> List[Dict]:
        """Convertit un tableau Word (lignes de cellules) en chunks via pipeline LLM."""
        return self._convert_table_rows(rows, heading)
    
    def chunk_pdf(self, file_path: Path) -> List[Dict]:
        """Chunking PDF par structure (TOC, font size, ou pages).
        
        Pipeline :
        1. Extraction content-based des tableaux (PyMuPDF find_tables → LLM)
        2. Chunking textuel (TOC / font / pages / vision)
        3. Fusion des deux flux
        
        Fallback vision : si le texte extractible est trop pauvre (infographie),
        convertit les pages en images et les passe à LLaVA pour description.
        """
        try:
            import fitz
            doc = fitz.open(file_path)
            
            # ── Passe 1 : Extraction content-based des tableaux ──
            table_chunks = self._extract_pdf_tables(doc, file_path.name)
            
            # ── Passe 2 : Chunking textuel (stratégies existantes) ──
            
            # Stratégie 1 : Table des matières
            toc = doc.get_toc()
            if toc and len(toc) > 2:
                text_chunks = self._chunk_pdf_by_toc(doc, toc)
                doc.close()
                if text_chunks:
                    return self._post_process(table_chunks + text_chunks)
            
            # Stratégie 2 : Détection titres par font size
            text_chunks = self._chunk_pdf_by_font(doc)
            doc.close()
            if text_chunks and len(text_chunks) > 1:
                return self._post_process(table_chunks + text_chunks)
            
            # Stratégie 3 : Par pages groupées intelligemment
            doc2 = fitz.open(file_path)
            text_pages = [doc2[i].get_text() for i in range(len(doc2))]
            n_pages = len(doc2)
            doc2.close()
            
            total_text = ''.join(t.strip() for t in text_pages)
            
            # Stratégie 4 (FALLBACK VISION) : PDF infographique
            # Si très peu de texte extractible → contenu visuel, passer à LLaVA
            if len(total_text) < 500 and n_pages <= 10:
                vision_chunks = self._chunk_pdf_visual(file_path, n_pages)
                if vision_chunks:
                    return self._post_process(table_chunks + vision_chunks)
            
            if not any(t.strip() for t in text_pages) and not table_chunks:
                return []
            
            # Grouper pages pour atteindre ~target_size mots
            text_chunks = []
            current = {'text': '', 'heading': '', 'page_info': ''}
            page_start = 1
            
            for i, page_text in enumerate(text_pages):
                current['text'] += '\n' + page_text
                if not current['heading']:
                    current['heading'] = f'Pages {i+1}'
                    page_start = i + 1
                
                words = len(current['text'].split())
                if words >= self.target_size or i == len(text_pages) - 1:
                    if current['text'].strip():
                        page_end = i + 1
                        current['page_info'] = f'Page {page_start}' if page_start == page_end else f'Pages {page_start}-{page_end}'
                        text_chunks.append(current)
                    current = {'text': '', 'heading': '', 'page_info': ''}
                    page_start = i + 2
            
            return self._post_process(table_chunks + text_chunks)
        except Exception as e:
            logger.warning(f"PDF chunking failed for {file_path.name}: {e}")
            return []
    
    def _extract_pdf_tables(self, doc, filename: str) -> List[Dict]:
        """Extrait les tableaux d'un PDF via PyMuPDF find_tables() → pipeline LLM.
        
        Pour chaque page, détecte les tableaux, extrait les cellules comme
        listes de lignes, puis passe par _convert_table_rows() (même pipeline
        que HTML/DOCX/spreadsheets).
        """
        all_table_chunks = []
        
        for page_idx, page in enumerate(doc):
            try:
                tables = page.find_tables()
                if not tables.tables:
                    continue
                
                for table in tables.tables:
                    # Extraire les données : list[list[str|None]]
                    data = table.extract()
                    if not data or len(data) < 2:
                        continue
                    
                    # Normaliser : None → '', strip
                    rows = []
                    for row in data:
                        cells = [(c or '').strip() for c in row]
                        if any(c for c in cells):
                            rows.append(cells)
                    
                    if len(rows) < 2:
                        continue
                    
                    heading = f"Tableau page {page_idx + 1}"
                    table_chunks = self._convert_table_rows(rows, heading)
                    all_table_chunks.extend(table_chunks)
                    
            except Exception as e:
                logger.debug(f"Table extraction failed on {filename} p{page_idx+1}: {e}")
                continue
        
        if all_table_chunks:
            logger.info(f"📊 PDF {filename}: {len(all_table_chunks)} chunks tabulaires extraits")
        
        return all_table_chunks
    
    def _chunk_pdf_by_toc(self, doc, toc) -> List[Dict]:
        """Chunk PDF par table des matières"""
        chunks = []
        for i, (level, title, page_num) in enumerate(toc):
            start_page = page_num - 1
            end_page = toc[i+1][2] - 1 if i < len(toc) - 1 else len(doc)
            
            section_text = []
            for p in range(start_page, min(end_page, len(doc))):
                section_text.append(doc[p].get_text())
            
            text = '\n'.join(section_text)
            if text.strip():
                page_label = f"Page {page_num}" if start_page == end_page - 1 else f"Pages {page_num}-{end_page}"
                chunks.append({'text': text, 'heading': title, 'page_info': page_label})
        return chunks
    
    def _chunk_pdf_by_font(self, doc) -> List[Dict]:
        """Chunk PDF par détection de titres (font size)"""
        font_sizes = []
        blocks_data = []
        
        for page_idx, page in enumerate(doc):
            for block in page.get_text("dict")["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_text = ' '.join(s["text"] for s in line["spans"])
                        max_size = max((s["size"] for s in line["spans"]), default=0)
                        font_sizes.append(max_size)
                        blocks_data.append((line_text, max_size, page_idx + 1))
        
        if not font_sizes:
            return []
        
        median = sorted(font_sizes)[len(font_sizes) // 2]
        threshold = median * 1.2
        
        sections = []
        current = {'title': '', 'text': [], 'start_page': 1}
        
        for text, size, page_num in blocks_data:
            if size >= threshold and len(text.strip()) < 200 and text.strip():
                if current['text']:
                    sections.append(current)
                current = {'title': text.strip(), 'text': [], 'start_page': page_num}
            else:
                current['text'].append(text)
                current['end_page'] = page_num
        
        if current['text']:
            sections.append(current)
        
        result = []
        for s in sections:
            joined = '\n'.join(s['text'])
            if joined.strip():
                sp = s.get('start_page', 1)
                ep = s.get('end_page', sp)
                page_label = f"Page {sp}" if sp == ep else f"Pages {sp}-{ep}"
                result.append({'text': joined, 'heading': s['title'], 'page_info': page_label})
        return result
    
    def _chunk_pdf_visual(self, file_path: Path, n_pages: int) -> List[Dict]:
        """Fallback vision pour PDFs infographiques.
        
        Convertit chaque page en image et la passe à LLaVA pour obtenir
        une description textuelle riche du contenu visuel.
        Utilisé quand PyMuPDF n'extrait presque pas de texte.
        """
        try:
            import fitz
            import base64
            import io
            import requests
            
            OLLAMA_URL = "http://localhost:11434"
            VISION_MODEL = "llava:7b"
            MAX_DIM = 1024  # Redimensionner pour VRAM
            
            VISION_PROMPT = """Tu analyses une page d'un document PDF de la CNIL (Commission Nationale Informatique et Libertés).
Cette page contient une infographie, un schéma ou un graphique avec peu de texte extractible.

Décris en détail et en français :
1. Le CONTENU PRINCIPAL : que montre cette page ? (données, processus, statistiques, schéma...)
2. Le TEXTE VISIBLE : retranscris tous les textes, titres, légendes, chiffres visibles
3. La STRUCTURE : est-ce un schéma de processus, un graphique à barres, un tableau, une infographie ?
4. L'UTILITÉ DPO : en quoi ce contenu est pertinent pour un Délégué à la Protection des Données ?

Sois exhaustif sur les données chiffrées et le texte visible. Pas de préambule, va droit au contenu."""
            
            chunks = []
            doc = fitz.open(file_path)
            
            for page_idx in range(min(n_pages, 10)):  # Max 10 pages
                page = doc[page_idx]
                
                # Convertir page en image (pixmap)
                # zoom=2 pour meilleure résolution → LLaVA voit mieux les textes
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("jpeg")
                
                # Redimensionner si trop grand
                from PIL import Image
                img = Image.open(io.BytesIO(img_bytes))
                w, h = img.size
                if max(w, h) > MAX_DIM:
                    ratio = MAX_DIM / max(w, h)
                    img = img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
                
                buf = io.BytesIO()
                img.save(buf, format='JPEG', quality=85)
                image_b64 = base64.b64encode(buf.getvalue()).decode('utf-8')
                
                # Appel LLaVA
                try:
                    response = requests.post(
                        f'{OLLAMA_URL}/api/generate',
                        json={
                            'model': VISION_MODEL,
                            'prompt': VISION_PROMPT,
                            'images': [image_b64],
                            'stream': False,
                            'options': {
                                'temperature': 0.1,
                                'num_predict': 500,
                            }
                        },
                        timeout=120
                    )
                    
                    if response.status_code != 200:
                        logger.warning(f"LLaVA HTTP {response.status_code} pour {file_path.name} p{page_idx+1}")
                        continue
                    
                    description = response.json().get('response', '').strip()
                    
                    if description and len(description) > 50:
                        # Récupérer aussi le peu de texte PyMuPDF comme complément
                        pdf_text = page.get_text().strip()
                        
                        text_parts = [f"[Infographie PDF - Page {page_idx+1}]"]
                        text_parts.append(f"Description visuelle : {description}")
                        if pdf_text and len(pdf_text) > 20:
                            text_parts.append(f"Texte extrait : {pdf_text}")
                        
                        chunks.append({
                            'text': '\n'.join(text_parts),
                            'heading': f"Infographie page {page_idx+1}",
                            'page_info': f"Page {page_idx+1}",
                        })
                        logger.info(f"📸 Vision PDF {file_path.name} p{page_idx+1}: {len(description)} chars")
                    else:
                        logger.debug(f"LLaVA description trop courte pour {file_path.name} p{page_idx+1}")
                        
                except requests.exceptions.Timeout:
                    logger.warning(f"LLaVA timeout pour {file_path.name} p{page_idx+1}")
                except Exception as e:
                    logger.warning(f"LLaVA erreur pour {file_path.name} p{page_idx+1}: {e}")
            
            doc.close()
            return chunks
            
        except Exception as e:
            logger.warning(f"PDF vision fallback failed for {file_path.name}: {e}")
            return []
    
    def chunk_doc(self, file_path: Path, file_type: str) -> List[Dict]:
        """Chunking documents Office (xlsx, ods, docx, odt)"""
        try:
            if file_type in ['xlsx', 'ods']:
                return self._chunk_spreadsheet(file_path, file_type)
            elif file_type in ['docx', 'odt']:
                return self._chunk_word(file_path, file_type)
            else:
                return []
        except Exception as e:
            logger.warning(f"Doc chunking failed for {file_path.name}: {e}")
            return []
    
    def _chunk_spreadsheet(self, file_path: Path, file_type: str) -> List[Dict]:
        """Chunking spreadsheets sémantique v2 : segmentation en zones + LLM Nemo.
        
        Stratégie :
        1. Extraire toutes les lignes comme listes de cellules
        2. Segmenter chaque feuille en zones sémantiques (séparées par lignes vides)
        3. Sérialiser chaque zone en texte tabulaire lisible
        4. Appeler LLM Nemo pour convertir en texte naturel (phrases sujet-verbe-complément)
        5. Découper si trop long
        """
        MAX_SHEETS = 20
        MAX_ROWS = 500
        
        all_sheets_rows = self._extract_spreadsheet_rows(file_path, file_type, MAX_SHEETS, MAX_ROWS)
        
        # Initialiser LLM si pas encore fait
        if not hasattr(self, '_llm'):
            try:
                config = RAGConfig()
                self._llm = config.llm_provider
            except Exception as e:
                logger.warning(f"LLM indisponible pour conversion tableur: {e}")
                self._llm = None
        
        chunks = []
        for sheet_name, rows in all_sheets_rows:
            zones = self._segment_sheet_zones(rows)
            
            for zone_idx, zone_rows in enumerate(zones):
                # Découper les zones trop grandes (taille dynamique ~250 mots)
                sub_zones = self._split_large_zone(zone_rows)
                
                for sub_zone in sub_zones:
                    # Sérialiser la zone en texte tabulaire brut
                    raw_table_text = self._zone_to_table_text(sub_zone, sheet_name)
                    
                    if not raw_table_text or len(raw_table_text.split()) < 10:
                        continue
                    
                    # Conversion LLM → texte naturel + tags
                    llm_result = self._llm_convert_table(raw_table_text, sheet_name)
                    natural_text = llm_result['text']
                    tags = llm_result.get('tags', [])
                    
                    if not natural_text or len(natural_text.split()) < 15:
                        continue
                    
                    # Découper si trop long (>500 mots → split par paragraphes)
                    sub_chunks = self._split_long_text(natural_text, sheet_name, tags=tags, max_words=500)
                    chunks.extend(sub_chunks)
        
        return self._post_process(chunks)
    
    def _extract_spreadsheet_rows(self, file_path: Path, file_type: str,
                                   max_sheets: int, max_rows: int) -> List[tuple]:
        """Extrait les lignes brutes de chaque feuille comme listes de cellules.
        
        Returns:
            [(sheet_name, [row1_cells, row2_cells, ...]), ...]
            Chaque row_cells est une liste de strings (cellules non vides + position).
        """
        result = []
        
        if file_type == 'xlsx':
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames[:max_sheets]:
                sheet = wb[sheet_name]
                rows = []
                for row in list(sheet.iter_rows(values_only=True))[:max_rows]:
                    cells = [str(c).strip() if c is not None else '' for c in row]
                    # Garder la ligne si au moins une cellule non vide
                    if any(c for c in cells):
                        rows.append(cells)
                if rows:
                    result.append((sheet_name, rows))
            wb.close()
        
        elif file_type == 'ods':
            from odf import table as odf_table, text as odf_text
            from odf.opendocument import load
            doc = load(str(file_path))
            sheets = doc.spreadsheet.getElementsByType(odf_table.Table)
            for i, sheet in enumerate(sheets[:max_sheets]):
                sheet_name = sheet.getAttribute('name') or f"Sheet{i+1}"
                rows = []
                for row in sheet.getElementsByType(odf_table.TableRow)[:max_rows]:
                    cell_elements = row.getElementsByType(odf_table.TableCell)
                    cells = []
                    for cell in cell_elements:
                        paras = cell.getElementsByType(odf_text.P)
                        cell_text = ' '.join(str(p) for p in paras).strip()
                        cells.append(cell_text)
                    if any(c for c in cells):
                        rows.append(cells)
                if rows:
                    result.append((sheet_name, rows))
        
        return result
    
    def _segment_sheet_zones(self, rows: List[List[str]]) -> List[List[List[str]]]:
        """Segmente une feuille en zones sémantiques séparées par des lignes vides.
        
        Chaque zone est un bloc contigu de lignes non-vides.
        Les zones trop petites (<2 lignes) sont fusionnées avec la suivante.
        """
        zones = []
        current_zone = []
        
        for row in rows:
            non_empty = [c for c in row if c.strip()]
            if not non_empty:
                # Ligne vide → fin de zone
                if current_zone:
                    zones.append(current_zone)
                    current_zone = []
            else:
                current_zone.append(row)
        
        # Dernière zone
        if current_zone:
            zones.append(current_zone)
        
        # Fusionner zones trop petites (1 ligne) avec la suivante
        merged = []
        buffer = None
        for zone in zones:
            if buffer is not None:
                # Fusionner le buffer avec cette zone
                merged.append(buffer + zone)
                buffer = None
            elif len(zone) == 1 and len(zone[0]) > 0:
                # Zone d'1 ligne : potentiellement un titre, buffer pour fusionner
                # Sauf si la cellule est très longue (>200 chars) → c'est un bloc autonome
                max_cell = max(len(c) for c in zone[0] if c.strip()) if any(c.strip() for c in zone[0]) else 0
                if max_cell > 200:
                    merged.append(zone)
                else:
                    buffer = zone
            else:
                merged.append(zone)
        
        if buffer:
            if merged:
                merged[-1] = merged[-1] + buffer
            else:
                merged.append(buffer)
        
        return merged
    
    def _split_large_zone(self, zone_rows: List[List[str]], max_words: int = 250) -> List[List[List[str]]]:
        """Découpe une zone trop grande en sous-zones.
        
        Adapte dynamiquement le nombre de lignes par sous-zone pour
        que chaque sous-zone produise ~max_words mots en entrée LLM.
        La ligne d'en-têtes est préfixée à chaque sous-zone.
        """
        # Calculer la densité de mots par ligne
        total_words = sum(len(' '.join(c for c in row if c.strip()).split()) for row in zone_rows)
        avg_words_per_row = max(total_words / max(len(zone_rows), 1), 1)
        # max_rows dynamique : viser max_words mots, plancher 5, plafond 20
        max_rows = max(5, min(20, int(max_words / avg_words_per_row)))
        
        if len(zone_rows) <= max_rows:
            return [zone_rows]
        
        # Heuristique : la première ligne est-elle un en-tête ?
        # ≥2 cellules non vides, cellules courtes (<80 chars), pas numériques
        first_row = zone_rows[0]
        non_empty_first = [c for c in first_row if c.strip()]
        numeric_count = sum(1 for c in non_empty_first
                          if c.replace('.','').replace(',','').replace('-','').replace('/','').strip().isdigit())
        is_header = (
            len(non_empty_first) >= 2
            and all(len(c) < 80 for c in non_empty_first)
            and numeric_count < len(non_empty_first) * 0.5
        )
        
        header_row = zone_rows[0] if is_header else None
        data_start = 1 if is_header else 0
        data_rows = zone_rows[data_start:]
        
        sub_zones = []
        for i in range(0, len(data_rows), max_rows):
            batch = data_rows[i:i + max_rows]
            if header_row:
                sub_zones.append([header_row] + batch)
            else:
                sub_zones.append(batch)
        
        return sub_zones
    
    def _zone_to_table_text(self, zone_rows: List[List[str]], sheet_name: str) -> str:
        """Convertit une zone en texte tabulaire lisible pour le LLM.
        
        Format :
        - Feuille : « nom »
        - Ligne 1 : col0 | col1 | col2
        - Ligne 2 : col0 | col1 | col2
        """
        lines = [f"Feuille : « {sheet_name} »"]
        
        for row in zone_rows:
            # Filtrer les cellules vides en fin de ligne
            cells = [c.strip() for c in row]
            # Tronquer trailing empties
            while cells and not cells[-1]:
                cells.pop()
            if not cells:
                continue
            
            # Remplacer les vides intermédiaires par "(vide)"
            display_cells = [c if c else '(vide)' for c in cells]
            lines.append(' | '.join(display_cells))
        
        return '\n'.join(lines)
    
    def _llm_convert_table(self, raw_table_text: str, sheet_name: str) -> dict:
        """Appelle Nemo LLM pour convertir un texte tabulaire en langage naturel
        et extraire des tags de catégorisation RGPD.
        
        Returns:
            dict avec 'text' (str) et 'tags' (list[str])
        """
        PROMPT = """Tu es un assistant spécialisé en protection des données personnelles (RGPD/CNIL).

TÂCHE : Convertis ce tableau en texte naturel, lisible et complet, puis catégorise-le.

RÈGLES POUR LE TEXTE :
- Écris des phrases complètes sujet-verbe-complément
- Conserve TOUTES les informations (noms, dates, durées, chiffres, articles de loi)
- Ne résume pas, ne paraphrase pas : restitue fidèlement le contenu
- Identifie les en-têtes de colonnes ou de lignes et utilise-les comme contexte
- Si le tableau a des en-têtes en première ligne OU en première colonne, associe chaque valeur à son en-tête
- Pour les listes de valeurs, utilise des énumérations naturelles
- Ne mets pas de titre, écris directement le texte
- Écris en français (même si le tableau est en anglais, traduis)

RÈGLES POUR LES TAGS :
- À la toute fin, ajoute une ligne commençant par [TAGS] suivie des thèmes RGPD couverts
- Choisis 1 à 3 tags parmi : droits des personnes, consentement, sécurité des données, durée de conservation, sous-traitance, base légale, données sensibles, transfert hors UE, cookies, violation de données, transparence, DPO, vidéosurveillance, finalité du traitement, registre des traitements, AIPD, anonymisation, minimisation, responsable de traitement, prospection commerciale, conformité RGPD, profilage, sanctions CNIL, données de santé, information des personnes
- Si aucun ne convient, propose un tag court et descriptif

TABLEAU :
{table_text}

TEXTE NATUREL :"""
        
        if self._llm is None:
            return self._mechanical_fallback(raw_table_text, sheet_name)
        
        try:
            prompt = PROMPT.format(table_text=raw_table_text)
            result = self._llm.generate(
                prompt,
                temperature=0.1,
                max_tokens=2000,
            )
            
            result = result.strip()
            
            if 'TABLEAU :' in result or 'RÈGLES POUR' in result:
                logger.warning(f"LLM a régurgité le prompt pour {sheet_name}, fallback")
                return self._mechanical_fallback(raw_table_text, sheet_name)
            
            # Détection copie brute : si le LLM a recopié le texte avec les pipes
            pipe_count = result.count(' | ')
            input_pipe_count = raw_table_text.count(' | ')
            if input_pipe_count > 3 and pipe_count > input_pipe_count * 0.5:
                logger.warning(f"LLM a recopié le tableau ({pipe_count} pipes) pour {sheet_name}, retry simplifié")
                # Retry avec un prompt plus directif
                retry_prompt = (
                    "Réécris ce contenu en français courant, sans utiliser de tableaux ni de tirets. "
                    "Chaque information doit être une phrase complète. "
                    "Conserve toutes les données (noms, dates, durées, chiffres).\n\n"
                    f"{raw_table_text}\n\nTexte réécrit :"
                )
                try:
                    result = self._llm.generate(retry_prompt, temperature=0.2, max_tokens=2000).strip()
                    pipe_count = result.count(' | ')
                    if pipe_count > input_pipe_count * 0.3:
                        logger.warning(f"Retry LLM toujours trop de pipes ({pipe_count}), fallback mécanique")
                        return self._mechanical_fallback(raw_table_text, sheet_name)
                except Exception as e:
                    logger.warning(f"Erreur retry LLM ({sheet_name}): {e}")
                    return self._mechanical_fallback(raw_table_text, sheet_name)
            
            # Extraire les tags [TAGS] de la fin de la réponse
            tags = []
            text_lines = result.split('\n')
            clean_lines = []
            for line in text_lines:
                if line.strip().startswith('[TAGS]') or line.strip().startswith('[Tags]'):
                    tags = parse_tags(line)
                else:
                    clean_lines.append(line)
            
            clean_text = '\n'.join(clean_lines).strip()
            return {
                'text': clean_text,
                'tags': tags,
            }
            
        except Exception as e:
            logger.warning(f"Erreur LLM conversion tableau ({sheet_name}): {e}")
            return self._mechanical_fallback(raw_table_text, sheet_name)
    
    def _mechanical_fallback(self, raw_table_text: str, sheet_name: str) -> dict:
        """Fallback mécanique si LLM indisponible : transformation propre sans pipes."""
        lines = raw_table_text.split('\n')
        text_parts = []
        for line in lines:
            if line.startswith('Feuille :'):
                continue
            # Séparer les cellules
            cells = [c.strip() for c in line.split(' | ')]
            cells = [c for c in cells if c and c != '(vide)']
            if not cells:
                continue
            if len(cells) == 1 and len(cells[0]) > 5:
                text_parts.append(cells[0])
            elif len(cells) == 2:
                # Format clé-valeur
                text_parts.append(f"{cells[0]} : {cells[1]}.")
            elif len(cells) >= 3:
                # Multi-colonnes : première cellule = label, reste = valeurs
                values = ', '.join(cells[1:])
                text_parts.append(f"{cells[0]} : {values}.")
        return {'text': '\n'.join(text_parts), 'tags': []}
    
    def _split_long_text(self, text: str, context_name: str, tags: List[str] = None,
                          max_words: int = 500) -> List[Dict]:
        """Découpe un texte trop long en chunks de taille raisonnable.
        Propage les tags RGPD dans chaque sous-chunk."""
        if tags is None:
            tags = []
        words = text.split()
        
        if len(words) <= max_words:
            return [{
                'text': text,
                'heading': context_name,
                'page_info': context_name,
                'rgpd_topics': tags,
            }]
        
        # Découper par paragraphes (double newline)
        paragraphs = text.split('\n\n')
        chunks = []
        current_text = []
        current_wc = 0
        
        for para in paragraphs:
            para_wc = len(para.split())
            if current_wc + para_wc > max_words and current_text:
                chunks.append({
                    'text': '\n\n'.join(current_text),
                    'heading': context_name,
                    'page_info': context_name,
                    'rgpd_topics': tags,
                })
                current_text = [para]
                current_wc = para_wc
            else:
                current_text.append(para)
                current_wc += para_wc
        
        if current_text:
            chunks.append({
                'text': '\n\n'.join(current_text),
                'heading': context_name,
                'page_info': context_name,
                'rgpd_topics': tags,
            })
        
        return chunks
    
    def _chunk_word(self, file_path: Path, file_type: str) -> List[Dict]:
        """Chunking documents texte par sections, avec détection de tableaux."""
        chunks = []
        
        if file_type == 'docx':
            from docx import Document
            from docx.oxml.ns import qn
            doc = Document(file_path)
            current = {'text': '', 'heading': ''}
            
            # Parcourir body dans l'ordre du DOM (paragraphes ET tables mélangés)
            body = doc.element.body
            for child in body:
                if child.tag == qn('w:p'):
                    # C'est un paragraphe
                    from docx.text.paragraph import Paragraph
                    para = Paragraph(child, body)
                    if para.style and para.style.name.startswith('Heading'):
                        if current['text'].strip():
                            chunks.append(current)
                        current = {'text': para.text, 'heading': para.text}
                    else:
                        current['text'] += '\n' + para.text
                elif child.tag == qn('w:tbl'):
                    # C'est un tableau → pipeline LLM
                    from docx.table import Table
                    table = Table(child, body)
                    rows = []
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells]
                        if any(c for c in cells):
                            rows.append(cells)
                    
                    if len(rows) >= 2:
                        # Flush texte courant
                        if current['text'].strip():
                            chunks.append(current)
                            current = {'text': '', 'heading': current.get('heading', '')}
                        
                        table_chunks = self._convert_docx_table(rows, current.get('heading', ''))
                        if table_chunks:
                            chunks.extend(table_chunks)
                    elif rows:
                        # Table triviale → texte plat
                        flat = ' '.join(' '.join(c for c in r if c) for r in rows)
                        current['text'] += '\n' + flat
            
            if current['text'].strip():
                chunks.append(current)
        
        elif file_type == 'odt':
            from odf import text as odf_text
            from odf.opendocument import load
            doc = load(str(file_path))
            all_text = []
            for h in doc.getElementsByType(odf_text.H):
                all_text.append(str(h))
            for p in doc.getElementsByType(odf_text.P):
                all_text.append(str(p))
            text = '\n'.join(all_text)
            if text.strip():
                chunks.append({'text': text, 'heading': file_path.stem})
        
        return self._post_process(chunks)
    
    def _post_process(self, chunks: List[Dict]) -> List[Dict]:
        """Post-traitement intelligent : merge tiny, split grands, overlap, heading prefix.
        
        Pipeline :
        1. Split les chunks >max_size aux frontières de phrases, avec overlap
        2. Fusionne les chunks <min_size avec leur voisin
        3. Préfixe le heading dans le texte (visible pour l'embedding)
        """
        if not chunks:
            return []
        
        # ── Étape 1 : Split des gros chunks aux frontières de phrases ──
        split_chunks = []
        for chunk in chunks:
            text = chunk['text'].strip()
            words = len(text.split())
            heading = chunk.get('heading', '')
            page_info = chunk.get('page_info', '')
            
            if words <= self.max_size:
                split_chunks.append({'text': text, 'heading': heading, 'page_info': page_info})
            else:
                # Split sémantique avec overlap
                sub_chunks = self._split_semantic(text, heading, page_info)
                split_chunks.extend(sub_chunks)
        
        # ── Étape 2 : Fusion des tiny chunks (<min_size) avec voisin ──
        merged = []
        for chunk in split_chunks:
            words = len(chunk['text'].split())
            if words < self.min_size and merged:
                # Fusionner avec le chunk précédent
                merged[-1]['text'] += ' ' + chunk['text']
                # Garder le heading du précédent sauf s'il est vide
                if not merged[-1]['heading'] and chunk['heading']:
                    merged[-1]['heading'] = chunk['heading']
            else:
                merged.append(chunk)
        
        # Si le premier chunk est tiny, fusionner avec le suivant
        if len(merged) > 1 and len(merged[0]['text'].split()) < self.min_size:
            merged[1]['text'] = merged[0]['text'] + ' ' + merged[1]['text']
            if not merged[1]['heading'] and merged[0]['heading']:
                merged[1]['heading'] = merged[0]['heading']
            merged.pop(0)
        
        # Si le dernier chunk est tiny, fusionner avec l'avant-dernier
        if len(merged) > 1 and len(merged[-1]['text'].split()) < self.min_size:
            merged[-2]['text'] += ' ' + merged[-1]['text']
            merged.pop()
        
        # ── Étape 3 : Préfixer heading dans le texte (pour l'embedding) ──
        result = []
        for chunk in merged:
            heading = chunk.get('heading', '').strip()
            text = chunk['text'].strip()
            
            if heading and not text.startswith(heading):
                # Préfixer seulement si le texte ne commence pas déjà par le heading
                chunk['text'] = f"[{heading}] {text}"
            
            result.append(chunk)
        
        return result
    
    def _split_semantic(self, text: str, heading: str, page_info: str) -> List[Dict]:
        """Split un texte long aux frontières de phrases, avec overlap.
        
        Stratégie :
        1. Découper en phrases (sur \n\n, puis . en fallback)
        2. Accumuler des phrases jusqu'à ~target_size mots
        3. Quand on dépasse, créer un nouveau chunk
        4. Ajouter overlap mots du chunk précédent au début du suivant
        """
        # Découper en unités sémantiques (paragraphes d'abord, puis phrases)
        paragraphs = re.split(r'\n\s*\n', text)
        
        # Si un seul gros paragraphe, split sur les phrases
        sentences = []
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(para.split()) > self.target_size:
                # Découper le paragraphe en phrases
                for sent in re.split(r'(?<=[.!?;:])\s+', para):
                    sent = sent.strip()
                    if sent:
                        sentences.append(sent)
            else:
                sentences.append(para)
        
        if not sentences:
            return [{'text': text, 'heading': heading, 'page_info': page_info}]
        
        # Accumuler en chunks de ~target_size mots
        chunks = []
        current_sentences = []
        current_words = 0
        
        for sent in sentences:
            sent_words = len(sent.split())
            
            if current_words + sent_words > self.max_size and current_sentences:
                # Flush le chunk courant
                chunk_text = ' '.join(current_sentences)
                chunks.append({'text': chunk_text, 'heading': heading, 'page_info': page_info})
                
                # Overlap : garder les derniers ~overlap mots pour le prochain chunk
                if self.overlap > 0:
                    overlap_text = self._get_tail_words(chunk_text, self.overlap)
                    current_sentences = [overlap_text, sent]
                    current_words = len(overlap_text.split()) + sent_words
                else:
                    current_sentences = [sent]
                    current_words = sent_words
            else:
                current_sentences.append(sent)
                current_words += sent_words
        
        # Dernier chunk
        if current_sentences:
            chunk_text = ' '.join(current_sentences)
            chunks.append({'text': chunk_text, 'heading': heading, 'page_info': page_info})
        
        return chunks
    
    @staticmethod
    def _get_tail_words(text: str, n_words: int) -> str:
        """Retourne les n derniers mots d'un texte."""
        words = text.split()
        return ' '.join(words[-n_words:]) if len(words) > n_words else text


class ChunkClassifier:
    """Classification hybride avec test d'ambiguïté rigoureux"""
    
    PROMPT = """Ce chunk est ambigu entre :
- {option1}
- {option2}

Extrait :
{text}

Quelle est la nature principale ?
Réponds en JSON : {{"chunk_nature": "...", "confidence": 0.9}}"""
    
    def __init__(self, llm):
        self.llm = llm
        self.stats = {
            'heuristic_certain': 0,  # Règles dures (confidence ≥ 0.90)
            'heuristic_probable': 0,  # Heuristiques claires (≥ 0.70)
            'llm_ambiguous': 0,  # LLM appelé car ambigu
            'fallback': 0,  # Fallback document
        }
    
    def classify(self, chunk: Dict, features: Dict, doc_nature: str) -> Dict:
        """Classification avec vrai test d'ambiguïté"""
        
        # PHASE 1 : RÈGLES DURES (pas de LLM, confiance 100%)
        
        # Sanction = toujours certain
        if features['has_sanction']:
            self.stats['heuristic_certain'] += 1
            return {
                'chunk_nature': 'SANCTION',
                'confidence': 1.0,
                'method': 'rule_sanction',
                'sectors': features.get('sectors', [])
            }
        
        # Template/formulaire = toujours GUIDE
        if features['has_template']:
            self.stats['heuristic_certain'] += 1
            return {
                'chunk_nature': 'GUIDE',
                'confidence': 0.95,
                'method': 'rule_template',
                'sectors': features.get('sectors', [])
            }
        
        # Procédure claire (steps + liste) = toujours GUIDE
        if features['has_steps'] and features['has_numbered_list']:
            self.stats['heuristic_certain'] += 1
            return {
                'chunk_nature': 'GUIDE',
                'confidence': 0.95,
                'method': 'rule_procedure',
                'sectors': features.get('sectors', [])
            }
        
        # Article RGPD pur (sans steps) = DOCTRINE
        if features['has_article_rgpd'] and not features['has_steps'] and features['word_count'] > 100:
            self.stats['heuristic_certain'] += 1
            return {
                'chunk_nature': 'DOCTRINE',
                'confidence': 0.90,
                'method': 'rule_article_pure',
                'sectors': features.get('sectors', [])
            }
        
        # PHASE 2 : CALCUL SCORES (détection ambiguïté)
        
        scores = {
            'DOCTRINE': (
                features['doctrine_score'] +
                int(features['has_article_rgpd']) +
                int(features['has_legal_ref'])
            ),
            'GUIDE': (
                features['operational_score'] +
                int(features['has_steps']) +
                int(features.get('has_numbered_list', False))
            ),
            'TECHNIQUE': int(features['has_technical']) * 2,  # Poids fort
            'SANCTION': 0,  # Déjà traité
        }
        
        # Trier scores
        sorted_scores = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        top1_nature, top1_score = sorted_scores[0]
        top2_nature, top2_score = sorted_scores[1]
        
        # TEST D'AMBIGUÏTÉ ULTRA-STRICT (ajustement ChatGPT final)
        # Ambigu SEULEMENT si :
        # - Aucun signal (score = 0)
        # OU
        # - Égalité PARFAITE entre 2 candidats crédibles + chunk dense
        
        is_truly_ambiguous = (
            top1_score == 0 or
            (
                top2_score >= 2 and
                (top1_score - top2_score) == 0 and  # Égalité parfaite !
                features['word_count'] > 150  # Chunk dense
            )
        )
        
        # PHASE 3 : CLASSIFICATION FINALE
        
        # Si score clair, utiliser heuristique
        if not is_truly_ambiguous and top1_score >= 1:
            self.stats['heuristic_probable'] += 1
            return {
                'chunk_nature': top1_nature,
                'confidence': min(0.85, 0.55 + top1_score * 0.10),
                'method': f'heuristic_{top1_nature.lower()}',
                'sectors': features.get('sectors', [])
            }
        
        # MIXTE IMPLICITE (ajustement ChatGPT 2)
        # Si DOCTRINE + GUIDE crédibles mais chunk court → doc_nature (mixte implicite)
        if (
            scores['DOCTRINE'] >= 2 and
            scores['GUIDE'] >= 2 and
            features['word_count'] < 200
        ):
            self.stats['heuristic_probable'] += 1
            return {
                'chunk_nature': doc_nature,
                'confidence': 0.60,
                'method': 'implicit_mixte_short',
                'sectors': features.get('sectors', [])
            }
        
        # DOCTRINE PAR DÉFAUT (ajustement ChatGPT 3 : seuil abaissé 80 → 60)
        # Si chunk narratif sans signaux forts → DOCTRINE
        if (
            features['word_count'] >= 60 and  # Abaissé de 80 à 60
            not features['has_steps'] and
            not features['has_template'] and
            not features['has_sanction'] and
            not features['has_technical']
        ):
            self.stats['heuristic_probable'] += 1
            return {
                'chunk_nature': 'DOCTRINE',
                'confidence': 0.65,
                'method': 'default_doctrine_narrative',
                'sectors': features.get('sectors', [])
            }
        
        # Si vraiment ambigu ET assez long, appeler LLM
        if is_truly_ambiguous and features['word_count'] >= 100:
            return self._llm_classify_constrained(
                chunk, doc_nature, top1_nature, top2_nature, features
            )
        
        # Sinon fallback sur document
        self.stats['fallback'] += 1
        return {
            'chunk_nature': doc_nature,
            'confidence': 0.50,
            'method': 'fallback_document',
            'sectors': features.get('sectors', [])
        }
    
    def _llm_classify_constrained(self, chunk: Dict, doc_nature: str, 
                                   option1: str, option2: str, features: Dict) -> Dict:
        """LLM avec choix contraint entre 2 options"""
        
        self.stats['llm_ambiguous'] += 1
        
        try:
            text = chunk['text'][:600]  # Max 600 chars
            
            # Prompt CONTRAINT (pas de choix libre)
            prompt = self.PROMPT.format(
                option1=option1,
                option2=option2,
                text=text
            )
            
            response = self.llm.generate(prompt, temperature=0.0, max_tokens=50)
            response_clean = clean_llm_json_response(response)
            result = safe_parse_json(response_clean)
            
            chunk_nature = result.get('chunk_nature', doc_nature)
            confidence = result.get('confidence', 0.70)
            
            # VALIDATION : Si LLM répond hors des 2 options, forcer option1
            if chunk_nature not in [option1, option2]:
                chunk_nature = option1
                confidence = 0.60
            
            # Si confiance faible, fallback
            if confidence < 0.60:
                chunk_nature = doc_nature
                confidence = 0.50
            
            return {
                'chunk_nature': chunk_nature,
                'confidence': confidence,
                'method': 'llm_constrained',
                'sectors': features.get('sectors', [])
            }
        
        except Exception as e:
            logger.debug(f"Erreur LLM: {e}")
            self.stats['fallback'] += 1
            return {
                'chunk_nature': doc_nature,
                'confidence': 0.50,
                'method': 'llm_error_fallback',
                'sectors': features.get('sectors', [])
            }



class ChunkProcessor:
    """Processeur principal"""
    
    def __init__(self, project_root: str = '.'):
        self.project_root = Path(project_root)
        self.data_path = self.project_root / 'data'
        self.raw_cnil_path = self.data_path / 'raw' / 'cnil'
        self.keep_cnil_path = self.data_path / 'keep' / 'cnil'
        
        self.manifest_file = self.raw_cnil_path / 'keep_manifest.json'
        self.metadata_file = self.raw_cnil_path / 'document_metadata.json'
        self.output_file = self.raw_cnil_path / 'processed_chunks.jsonl'
        
        self.chunker = StructuralChunker()
        self.feature_extractor = ChunkFeatureExtractor()
        
        # LLM
        config = RAGConfig()
        self.llm = config.llm_provider
        self.classifier = ChunkClassifier(self.llm)
        
        # Flag d'interruption gracieuse
        self._interrupted = False
        self._original_sigint = signal.getsignal(signal.SIGINT)
        
        self.stats = {
            'docs': 0,
            'images': 0,
            'chunks': 0,
            'resumed_skip': 0,
            'by_nature': {'DOCTRINE': 0, 'GUIDE': 0, 'SANCTION': 0, 'TECHNIQUE': 0}
        }
        
        logger.info("🤖 Chunker initialisé")
    
    def _handle_interrupt(self, signum, frame):
        """Gestionnaire Ctrl+C : flag interruption pour flush propre."""
        if self._interrupted:
            print("\n\n⚠️  Deuxième Ctrl+C — arrêt immédiat !")
            sys.exit(1)
        self._interrupted = True
        print("\n\n🛑 Ctrl+C détecté — fin du document en cours puis arrêt...")
        print("   (Ctrl+C à nouveau pour arrêt immédiat)")
    
    def _load_already_chunked(self) -> set:
        """Charge les document_id déjà présents dans le JSONL existant."""
        done = set()
        if self.output_file.exists():
            try:
                with open(self.output_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            try:
                                chunk = json.loads(line)
                                doc_id = chunk.get('document_id', '')
                                if doc_id:
                                    done.add(doc_id)
                            except json.JSONDecodeError:
                                continue
                if done:
                    logger.info(f"♻️  JSONL existant : {len(done)} documents déjà chunkés")
            except Exception as e:
                logger.warning(f"⚠️  Impossible de lire le JSONL existant : {e}")
        return done
    
    def run(self, max_docs: int = None, fresh: bool = False):
        """Execute chunking avec support resume.
        
        Args:
            max_docs: Limiter à N documents (mode test)
            fresh: Si True, ignore le JSONL existant et recommence à zéro
        """
        
        print("=" * 70)
        print("PHASE 5B : CHUNKING INTELLIGENT (KEEP ONLY)")
        print("=" * 70)
        
        # Installer gestionnaire Ctrl+C
        signal.signal(signal.SIGINT, self._handle_interrupt)
        
        # Load keep_manifest.json (fichiers pertinents seulement)
        logger.info(f"📂 Chargement du manifest : {self.manifest_file}")
        with open(self.manifest_file, 'r', encoding='utf-8') as f:
            manifest_data = json.load(f)
        
        # Load document_metadata.json (classifications LLM)
        logger.info(f"📂 Chargement des classifications : {self.metadata_file}")
        with open(self.metadata_file, 'r', encoding='utf-8') as f:
            metadata_classifications = json.load(f).get('metadata', {})
        
        # Construire le cache URL depuis keep_manifest
        self.url_cache = {}  # file_path -> {url, parent_url}
        for list_key in ['html', 'pdfs', 'docs', 'images']:
            for item in manifest_data.get(list_key, []):
                fp = item.get('metadata', {}).get('file_path', '')
                if fp:
                    self.url_cache[fp] = {
                        'url': item.get('url', ''),
                        'parent_url': item.get('parent_url', item.get('metadata', {}).get('source_url', '') or ''),
                    }
        logger.info(f"📎 Cache URLs construit : {len(self.url_cache)} documents")
        
        # Construire la liste des documents à traiter depuis keep_manifest
        documents_to_process = []
        
        # HTML files
        for html_item in manifest_data.get('html', []):
            file_path = html_item['metadata']['file_path']
            doc_meta = metadata_classifications.get(file_path, {})
            if doc_meta:
                documents_to_process.append((file_path, doc_meta))
        
        # PDF files
        for pdf_item in manifest_data.get('pdfs', []):
            file_path = pdf_item['metadata']['file_path']
            doc_meta = metadata_classifications.get(file_path, {})
            if doc_meta:
                documents_to_process.append((file_path, doc_meta))
        
        # Doc files (xlsx, odt, docx)
        for doc_item in manifest_data.get('docs', []):
            file_path = doc_item['metadata']['file_path']
            doc_meta = metadata_classifications.get(file_path, {})
            if doc_meta:
                documents_to_process.append((file_path, doc_meta))
        
        # Images SCHEMA_DPO (classifiées par Phase 4B + 5A)
        image_classification_file = self.raw_cnil_path / 'image_classification.json'
        image_cls_data = {}
        if image_classification_file.exists():
            with open(image_classification_file, 'r', encoding='utf-8') as f:
                image_cls_data = json.load(f).get('classifications', {})
        
        for img_item in manifest_data.get('images', []):
            file_path = img_item['metadata']['file_path']
            doc_meta = metadata_classifications.get(file_path, {})
            if doc_meta:
                # Enrichir avec les données Phase 4B (OCR, LLaVA)
                img_name = Path(file_path).name
                img_4b = image_cls_data.get(img_name, {})
                doc_meta['_image_ocr_preview'] = img_4b.get('ocr_preview', '')
                doc_meta['_image_ocr_words'] = img_4b.get('ocr_words', 0)
                doc_meta['_image_description'] = img_4b.get('description', '')
                doc_meta['_image_raison'] = img_4b.get('raison', '')
                doc_meta['_image_parent_url'] = img_item.get('parent_url', '')
                doc_meta['_image_url'] = img_item.get('url', '')
                documents_to_process.append((file_path, doc_meta))
        
        if max_docs:
            documents_to_process = documents_to_process[:max_docs]
            print(f"\n  MODE TEST : {max_docs} documents")
        
        # Resume : charger docs déjà traités
        if fresh and self.output_file.exists():
            self.output_file.unlink()
            already_done = set()
            print(f"\n🗑️  Mode fresh : JSONL existant supprimé")
        else:
            already_done = self._load_already_chunked()
        
        total_docs = len(documents_to_process)
        new_docs = [d for d in documents_to_process if Path(d[0]).stem not in already_done]
        
        print(f"\n  Documents total     : {total_docs}")
        print(f"  Déjà chunkés (skip) : {total_docs - len(new_docs)}")
        print(f"  À chunker           : {len(new_docs)}")
        print(f"  Total classifiés    : {len(metadata_classifications)}")
        print(f"  Durée estimée       : ~{len(new_docs) * 3 / 60:.0f} minutes")
        
        if len(new_docs) == 0:
            print("\n✅ Tous les documents déjà chunkés !")
            signal.signal(signal.SIGINT, self._original_sigint)
            return
        
        if already_done:
            print(f"\n♻️  MODE RESUME : {len(already_done)} docs skippés, {len(new_docs)} restants")
            print(f"   (utiliser --fresh pour recommencer à zéro)")
        
        print(f"   💡 Ctrl+C pour interrompre proprement (reprise possible)\n")
        
        self.stats['resumed_skip'] = total_docs - len(new_docs)
        
        # Process — mode append pour le resume
        write_mode = 'a' if already_done else 'w'
        
        with open(self.output_file, write_mode, encoding='utf-8') as f_out:
            for file_path, doc_meta in tqdm(new_docs, desc="Chunking"):
                if self._interrupted:
                    break
                
                chunks = self._process_doc(file_path, doc_meta)
                
                for chunk in chunks:
                    f_out.write(json.dumps(chunk, ensure_ascii=False) + '\n')
                
                # Flush après chaque document pour ne rien perdre
                f_out.flush()
                
                self.stats['docs'] += 1
                self.stats['chunks'] += len(chunks)
        
        # Restaurer handler Ctrl+C
        signal.signal(signal.SIGINT, self._original_sigint)
        
        if self._interrupted:
            print(f"\n🛑 Interrompu après {self.stats['docs']} docs ({self.stats['chunks']} chunks).")
            print(f"   Relancez la commande pour reprendre automatiquement.")
        
        self._print_summary()
    
    def _process_doc(self, file_path: str, doc_meta: Dict) -> List[Dict]:
        """Process document"""
        
        fp = Path(file_path)
        if not fp.exists():
            return []
        
        # Chunking
        file_type = doc_meta.get('file_type', 'html')
        if file_type == 'html':
            raw_chunks = self.chunker.chunk_html(fp)
        elif file_type == 'pdf':
            raw_chunks = self.chunker.chunk_pdf(fp)
        elif file_type in ['xlsx', 'ods', 'docx', 'odt']:
            raw_chunks = self.chunker.chunk_doc(fp, file_type)
        elif file_type == 'image':
            raw_chunks = self._chunk_image(fp, doc_meta)
        else:
            logger.warning(f"Type non supporté: {file_type} pour {fp.name}")
            return []
        
        doc_nature = doc_meta.get('nature', 'GUIDE')
        doc_id = fp.stem
        
        # Classify chunks
        result_chunks = []
        
        for idx, raw_chunk in enumerate(raw_chunks):
            features = self.feature_extractor.extract(raw_chunk['text'])
            classification = self.classifier.classify(raw_chunk, features, doc_nature)
            
            chunk_nature = classification['chunk_nature']
            
            if chunk_nature in self.stats['by_nature']:
                self.stats['by_nature'][chunk_nature] += 1
            
            # Récupérer URLs et metadata du document
            url_info = self.url_cache.get(file_path, {})
            
            chunk = {
                'chunk_id': f"{doc_id}_{idx}",
                'document_id': doc_id,
                'document_path': file_path,
                'document_nature': doc_nature,
                'chunk_nature': chunk_nature,
                'chunk_index': self._map(chunk_nature),
                'sectors': classification.get('sectors', []),  # SECTEURS CHUNK-LEVEL
                'text': raw_chunk['text'][:5000],
                'heading': raw_chunk.get('heading', ''),
                'page_info': raw_chunk.get('page_info', ''),  # ex: "Pages 3-5", "Sheet: Registre"
                'confidence': classification['confidence'],
                'method': classification.get('method', 'unknown'),
                # Traçabilité source
                'file_type': file_type,
                'source_url': url_info.get('url', ''),
                'parent_url': url_info.get('parent_url', ''),
                'title': doc_meta.get('raison', '') or doc_meta.get('title', ''),  # raison = description LLM
            }
            
            result_chunks.append(chunk)
        
        return result_chunks
    
    def _chunk_image(self, file_path: Path, doc_meta: Dict) -> List[Dict]:
        """Crée un chunk textuel à partir d'une image SCHEMA_DPO.
        
        Combine : titre + description LLaVA + texte OCR + lien parent.
        Produit 1 chunk par image (les schémas sont des unités atomiques).
        """
        ocr_text = doc_meta.get('_image_ocr_preview', '')
        ocr_words = doc_meta.get('_image_ocr_words', 0)
        description = doc_meta.get('_image_description', '')
        raison = doc_meta.get('_image_raison', '')
        parent_url = doc_meta.get('_image_parent_url', '')
        image_url = doc_meta.get('_image_url', '')
        title = doc_meta.get('title', '') or doc_meta.get('raison', '') or f"Schéma DPO - {file_path.stem}"
        
        # Construire le texte du chunk : description + OCR
        text_parts = []
        
        text_parts.append(f"[Schéma/Diagramme DPO] {title}")
        
        if description and description != title:
            text_parts.append(f"Description visuelle : {description}")
        
        if raison:
            text_parts.append(f"Pertinence DPO : {raison}")
        
        if parent_url:
            text_parts.append(f"Page source : {parent_url}")
        
        if ocr_text:
            # Nettoyer l'OCR (souvent bruité sur les schémas)
            clean_ocr = ' '.join(ocr_text.split())  # Normaliser espaces
            text_parts.append(f"Contenu textuel extrait (OCR) : {clean_ocr}")
        
        text = '\n'.join(text_parts)
        
        if len(text.split()) < 10:
            logger.debug(f"Image {file_path.name} : contenu trop pauvre ({len(text.split())} mots), skip")
            return []
        
        return [{
            'text': text,
            'heading': title,
            'page_info': f"Image: {image_url}" if image_url else f"Image: {file_path.name}",
        }]
    
    def _map(self, nature: str) -> str:
        """Map nature → index"""
        mapping = {
            'DOCTRINE': 'FONDAMENTAUX',
            'GUIDE': 'OPERATIONNEL',
            'SANCTION': 'SANCTIONS',
            'TECHNIQUE': 'TECHNIQUE',
        }
        return mapping.get(nature, 'OPERATIONNEL')
    
    def _print_summary(self):
        """Summary"""
        print("\n" + "=" * 70)
        print("📊 RÉSUMÉ")
        print("=" * 70)
        if self.stats['resumed_skip'] > 0:
            print(f"\n♻️  Repris    : {self.stats['resumed_skip']} docs déjà traités (skippés)")
        print(f"\n📄 Documents  : {self.stats['docs']} (nouvellement traités)")
        print(f"📦 Chunks     : {self.stats['chunks']}")
        print(f"📊 Moyenne    : {self.stats['chunks'] / max(1, self.stats['docs']):.1f} chunks/doc")
        
        print(f"\n🔬 Distribution par NATURE :")
        for nature, count in sorted(self.stats['by_nature'].items()):
            pct = count / max(1, self.stats['chunks']) * 100
            print(f"   {nature:12s} : {count:6d} ({pct:5.1f}%)")
        
        print(f"\n⚙️  Méthodes de classification :")
        total_methods = sum(self.classifier.stats.values())
        for method, count in sorted(self.classifier.stats.items()):
            pct = count / max(1, total_methods) * 100
            print(f"   {method:25s} : {count:6d} ({pct:5.1f}%)")
        
        # Calculer % LLM
        llm_pct = self.classifier.stats.get('llm_ambiguous', 0) / max(1, self.stats['chunks']) * 100
        print(f"\n💡 Appels LLM : {llm_pct:.1f}% des chunks (cible < 10%)")
        
        print("\n" + "=" * 70)
        print(f"💾 Output : {self.output_file}")
        print("=" * 70)


def main():
    import argparse
    parser = argparse.ArgumentParser(description='Phase 5B : Chunking intelligent')
    parser.add_argument('--test', type=int, help='Limiter à N documents')
    parser.add_argument('--fresh', action='store_true', help='Ignorer le JSONL existant, recommencer à zéro')
    args = parser.parse_args()
    
    processor = ChunkProcessor()
    processor.run(max_docs=args.test, fresh=args.fresh)


if __name__ == "__main__":
    main()